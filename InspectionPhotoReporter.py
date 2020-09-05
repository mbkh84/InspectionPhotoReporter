from tkinter import Tk, Label, Button, Canvas, Radiobutton, filedialog
from tkinter import *
from PIL import Image
from tkinter.ttk import Progressbar
import os
import docx
from docx.shared import Inches
from docx.shared import Mm


class MyFirstGUI:
    def __init__(self, master):
        self.master = master
        master.title("Photo Docx Report Generator")
        self.m_selectedPath=StringVar()
        self.m_pixWidth=StringVar()
        self.m_pixHeight=StringVar()
        self.m_rprtCol=StringVar()
        self.m_imgSize =StringVar()

        self.myCanvas=Canvas(master,height=270,width=325)
        self.myCanvas.pack()

        self.label = Label(master, text="Folder of Photos:")
        self.label.place(x=10,y=10)

        self.addrLabel = Label(master, text='No address is selected ')
        self.addrLabel.place(x=110,y=10)

        self.search_dir_button=Button(master, text="Serach", command=self.choose_folder)
        self.search_dir_button.place(x=10,y=40)

        # global m_pixWidth
        m_tPosy=70
        m_tPosx=10
        # m_pixWidth=StringVar()
        self.label = Label(master, text="New Picture Width (pxl):")
        self.label.place(x=m_tPosx,y=m_tPosy)
        self.m_entEff = Entry(master,  width=20,textvariable=self.m_pixWidth)
        self.m_entEff.place(x=m_tPosx+170,y=m_tPosy)

        # global m_pixHeight
        m_tPosy=100
        m_tPosx=10
        # m_pixHeight=StringVar()
        self.label = Label(master, text="New Picture Height (pxl):")
        self.label.place(x=m_tPosx,y=m_tPosy)
        self.m_entEff = Entry(master,  width=20,textvariable=self.m_pixHeight)
        self.m_entEff.place(x=m_tPosx+170,y=m_tPosy)

        # global m_rprtCol
        m_tPosy=130
        m_tPosx=10
        # m_rprtCol=StringVar()
        self.label = Label(master, text="Number of Columns")
        self.label.place(x=m_tPosx,y=m_tPosy)
        self.m_entEff = Entry(master,  width=20,textvariable=self.m_rprtCol)
        self.m_entEff.place(x=m_tPosx+170,y=m_tPosy)

        # global m_imgSize
        m_tPosy=160
        m_tPosx=10
        # m_imgSize=StringVar()
        self.label = Label(master, text="Image Size in Table Cell(mm)")
        self.label.place(x=m_tPosx,y=m_tPosy)
        self.m_entEff = Entry(master,  width=20,textvariable=self.m_imgSize)
        self.m_entEff.place(x=m_tPosx+170,y=m_tPosy)

        self.gen_button=Button(master, text="Generate Report", command=self.gen_report)
        self.gen_button.place(x=10,y=190)

        self.m_procProg=Progressbar(master,orient=HORIZONTAL,length=300,mode='determinate')
        self.m_procProg.place(x=10,y=225)


    def choose_folder(self):
        m_temp= filedialog.askdirectory()
        self.m_selectedPath.set(m_temp)
        self.addrLabel.configure(text = m_temp)

    def gen_report(self):
        errorString=''
        errorFlag=False
        if( not self.m_selectedPath.get()):
            errorString = ' Please Input a Valid Folder \n'
            errorFlag = True
        if( not self.m_pixWidth.get()):
            errorString += ' Please Input a Valid Number for Photo Width \n'
            errorFlag = True
        else:
            m_pixelWidth=int(self.m_pixWidth.get())
            if(not (m_pixelWidth>=50 and m_pixelWidth<=3000)):
                errorString += ' Please Input a Valid Number for Photo Width 50-3000 \n'
                errorFlag = True
        if( not self.m_pixHeight.get()):
            errorString += ' Please Input a Valid Number for Photo Height \n'
            errorFlag = True
        else:
            m_pixelHeight=int(self.m_pixHeight.get())
            if(not (m_pixelHeight>=50 and m_pixelHeight<=3000)):
                errorString += ' Please Input a Valid Number for Photo Heigth 50-3000 \n'
                errorFlag = True
        if( not self.m_rprtCol.get()):
            errorString += ' Please Input a Valid Number for the number of columns \n'
            errorFlag = True
        else:
            m_rprtCol=int(self.m_rprtCol.get())
            if(not (m_rprtCol>=1 and m_rprtCol<=10)):
                errorString += ' Please Input a Valid Number for Columns 1-10 \n'
                errorFlag = True
        if( not self.m_imgSize.get()):
            errorString += ' Please Input a Valid Number for image size \n'
            errorFlag = True
        else:
            m_imgSize=int(self.m_imgSize.get())
            if(not (m_imgSize>=10 and m_imgSize<=210)):
                errorString += ' Please Input a Valid Number for Columns 10-210 \n'
                errorFlag = True

        if(errorFlag):
            self.popupmsg(errorString)

        else:
            size_image = (m_pixelWidth,m_pixelHeight)
            m_parent=self.m_selectedPath.get()
            m_dir='diffSize'
            m_path = m_parent+'/'+ m_dir
            print (m_path)
            if not os.path.exists(m_path):
                os.mkdir(m_path)
            else:
                print('The Folder Was Avilable Before!')
                for f1 in os.listdir(m_path):
                    if f1.endswith('.jpg'):
                        os.remove(m_path+'/'+f1)
                print('The Folder Was Cleared!')

            m_numFiles=0;
            for f in os.listdir(m_parent):
                if f.endswith('.jpg'):
                    m_numFiles = m_numFiles + 1
            print(m_numFiles)
            self.m_procProg["maximum"]=m_numFiles

            m_col=m_rprtCol
            m_imageLengthInWord=m_imgSize
            m_i=0
            m_j=0
            m_row=int(m_numFiles/m_col) +1
            m_row=m_row * 2
            print (m_row , m_col)
            document = docx.Document()

            section = document.sections[0]
            section.page_height = Mm(297)
            section.page_width = Mm(210)
            section.left_margin = Mm(25.4)
            section.right_margin = Mm(25.4)
            section.top_margin = Mm(25.4)
            section.bottom_margin = Mm(15.4)
            section.header_distance = Mm(12.7)
            section.footer_distance = Mm(12.7)

            tbl = document.add_table(rows=m_row, cols=m_col)
            tbl.style = 'TableGrid'
            tbl.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
            # tbl.style = 'LightShading-Accent'

            m_tempCounter=0

            for f in os.listdir(m_parent):
                if f.endswith('.jpg'):
                    self.m_procProg["value"]=m_tempCounter
                    self.m_procProg.update()
                    m_tempCounter = m_tempCounter +1
                    i=Image.open(m_parent+'/'+f)
                    fn,fext = os.path.splitext(f)
                    width, height = i.size
                    if height>width:
                        i=i.transpose(Image.ROTATE_90)
                    i.thumbnail(size_image)
                    i.save(m_path+'/{}_300{}'.format(fn,fext))

                    cell = tbl.cell(m_j,m_i)
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(m_path+'/'+fn+'_300'+fext, width=Mm(m_imageLengthInWord))

                    cell = tbl.cell(m_j+1,m_i)
                    cell.text='Fig-'+str(m_j*m_col + (m_i+1))

                    if(m_i +1 ==m_col):
                        m_i=0
                        m_j = m_j +2
                    else:
                        m_i = m_i + 1
            document.save(m_path+'/'+'PhotoPortfo.docx')

            self.popupmsg('Process Successfully Finished!')


    def popupmsg(self,msg):
        popup=Tk()
        popup.wm_title('!')
        label = Label(popup,text=msg)
        label.pack()#side='top' , fill='x' , pady=1)

root = Tk()
my_gui = MyFirstGUI(root)
root.mainloop()
